#!/usr/bin/env python3
"""
md_to_docx_resume.py
Converts resume-source.md into jennifer.docx using python-docx.

Usage:
    python md_to_docx_resume.py                          # reads resume-source.md, writes jennifer.docx
    python md_to_docx_resume.py --input my.md --output out.docx

The Markdown file is the single source of truth.
Edit the Markdown, run this script to generate the Word document.
"""

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Reuse the parser from the LaTeX converter
from md_to_latex_resume import parse_markdown


# -- Style constants --
FONT_NAME = "Calibri"
COLOR_ACCENT = RGBColor(0xD0, 0x6B, 0x00)  # warm orange, matching moderncv
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_GRAY = RGBColor(0x55, 0x55, 0x55)


def set_run_font(run, size=10, bold=False, color=COLOR_BLACK, name=FONT_NAME):
    """Apply font settings to a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_section_heading(doc, text):
    """Add an orange section heading with a bottom rule."""
    para = doc.add_paragraph()
    para.space_before = Pt(8)
    para.space_after = Pt(2)
    run = para.add_run(text.upper())
    set_run_font(run, size=11, bold=True, color=COLOR_ACCENT)
    # Thin bottom border via paragraph border XML
    from docx.oxml.ns import qn
    from lxml import etree
    pPr = para._p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D06B00")


def add_bullet(doc, text, indent=0.25):
    """Add a bullet-pointed paragraph."""
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(indent)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    run = para.add_run(text)
    set_run_font(run, size=9.5)
    return para


def add_entry_header(doc, left_text, right_text=""):
    """Add a bold entry header with optional right-aligned date."""
    para = doc.add_paragraph()
    para.space_before = Pt(6)
    para.space_after = Pt(1)
    run = para.add_run(left_text)
    set_run_font(run, size=10, bold=True)
    if right_text:
        run2 = para.add_run(f"\t{right_text}")
        set_run_font(run2, size=9.5, color=COLOR_GRAY)
        # Right-align tab stop
        from docx.oxml.ns import qn
        from lxml import etree
        pPr = para._p.get_or_add_pPr()
        tabs = etree.SubElement(pPr, qn("w:tabs"))
        tab = etree.SubElement(tabs, qn("w:tab"))
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "9360")  # ~6.5 inches in twips
    return para


def add_detail_line(doc, text, size=9, color=COLOR_GRAY):
    """Add a smaller gray detail line (location, clearance, etc.)."""
    para = doc.add_paragraph()
    para.space_before = Pt(0)
    para.space_after = Pt(1)
    run = para.add_run(text)
    set_run_font(run, size=size, color=color)
    return para


def generate_docx(resume: dict, output_path: str):
    """Generate a Word document from parsed resume data."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # -- Name --
    name = f"{resume['name_first']} {resume['name_last']}"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.space_after = Pt(2)
    run = para.add_run(name)
    set_run_font(run, size=20, bold=True, color=COLOR_ACCENT)

    # -- Contact info --
    contact_parts = []
    if resume["location"]:
        contact_parts.append(resume["location"])
    if resume.get("remote"):
        contact_parts.append(f"Remote, {resume['remote']}")
    if resume["email"]:
        contact_parts.append(resume["email"])
    if resume["linkedin"]:
        contact_parts.append(f"linkedin.com/in/{resume['linkedin']}")
    if resume["phone"]:
        contact_parts.append(resume["phone"])

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.space_before = Pt(0)
    para.space_after = Pt(4)
    run = para.add_run("  |  ".join(contact_parts))
    set_run_font(run, size=9, color=COLOR_GRAY)

    # -- Summary --
    if resume["summary"]:
        para = doc.add_paragraph()
        para.space_before = Pt(2)
        para.space_after = Pt(2)
        parts = [p.strip() for p in resume["summary"].split("|")]
        run = para.add_run(" | ".join(parts))
        set_run_font(run, size=9.5)

    # -- Credentials --
    if resume["credentials"]:
        para = doc.add_paragraph()
        para.space_before = Pt(0)
        para.space_after = Pt(2)
        run = para.add_run("Clearance & Credentials: ")
        set_run_font(run, size=9.5, bold=True)
        run2 = para.add_run(resume["credentials"])
        set_run_font(run2, size=9.5)

    # -- Education summary --
    if resume["education_summary"]:
        para = doc.add_paragraph()
        para.space_before = Pt(0)
        para.space_after = Pt(2)
        run = para.add_run("Education: ")
        set_run_font(run, size=9.5, bold=True)
        run2 = para.add_run(resume["education_summary"])
        set_run_font(run2, size=9.5)

    # -- Technical Proficiency --
    if resume["technical"]:
        add_section_heading(doc, "Technical Proficiency")
        for category, items in resume["technical"]:
            para = doc.add_paragraph()
            para.space_before = Pt(1)
            para.space_after = Pt(1)
            run = para.add_run(f"{category}: ")
            set_run_font(run, size=9.5, bold=True)
            run2 = para.add_run(items)
            set_run_font(run2, size=9.5)

    # -- Publication --
    if resume["publication"]:
        add_section_heading(doc, "Publication")
        pub = resume["publication"]
        # Strip markdown bold markers
        pub_clean = pub.replace("**", "")
        para = doc.add_paragraph()
        para.space_before = Pt(2)
        para.space_after = Pt(2)
        run = para.add_run(pub_clean)
        set_run_font(run, size=9.5)

    # -- Experience --
    if resume["experience"]:
        add_section_heading(doc, "Experience")
        for entry in resume["experience"]:
            header = f"{entry['title']}  |  {entry['company']}"
            add_entry_header(doc, header, entry["dates"])
            if entry["location"]:
                add_detail_line(doc, entry["location"])
            for bullet in entry["bullets"]:
                add_bullet(doc, bullet)

    # -- Notable Projects --
    if resume["notable_projects"]:
        add_section_heading(doc, "Notable Projects")
        for subsection, projects in resume["notable_projects"].items():
            para = doc.add_paragraph()
            para.space_before = Pt(6)
            para.space_after = Pt(2)
            run = para.add_run(subsection)
            set_run_font(run, size=10, bold=True, color=COLOR_ACCENT)

            for project in projects:
                header = f"{project['year']}  |  {project['title']}"
                add_entry_header(doc, header)
                for bullet in project["bullets"]:
                    add_bullet(doc, bullet)

    # -- References --
    if resume["references"]:
        add_section_heading(doc, "References")
        for ref in resume["references"]:
            ref_clean = ref.replace("**", "")
            parts = [p.strip() for p in ref_clean.split("|")]
            para = doc.add_paragraph()
            para.space_before = Pt(3)
            para.space_after = Pt(1)
            if len(parts) >= 5:
                run = para.add_run(parts[0])
                set_run_font(run, size=10, bold=True)
                detail = f"  |  {parts[1]}  |  {parts[2]}"
                run2 = para.add_run(detail)
                set_run_font(run2, size=9.5)
                contact_para = doc.add_paragraph()
                contact_para.space_before = Pt(0)
                contact_para.space_after = Pt(1)
                run3 = contact_para.add_run(f"{parts[3]}  |  {parts[4]}")
                set_run_font(run3, size=9, color=COLOR_GRAY)
            else:
                run = para.add_run(ref_clean)
                set_run_font(run, size=9.5)

    # -- Education Detail --
    if resume["education_detail"]:
        add_section_heading(doc, "Education")
        for edu in resume["education_detail"]:
            header = edu["header"].replace("**", "")
            parts = [p.strip() for p in header.split("|")]
            if len(parts) >= 3:
                add_entry_header(doc, f"{parts[1]}  |  {parts[2]}", parts[0])
            else:
                add_entry_header(doc, header)
            for bullet in edu["bullets"]:
                add_bullet(doc, bullet)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(
        description="Convert resume Markdown to Word (.docx)"
    )
    parser.add_argument(
        "--input", default="resume-source.md", help="Input Markdown file"
    )
    parser.add_argument(
        "--output", default="jennifer.docx", help="Output Word file"
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: {input_path} not found", file=sys.stderr)
        sys.exit(1)

    resume = parse_markdown(str(input_path))
    generate_docx(resume, args.output)
    print(f"Generated {args.output} from {input_path}")
    print(f"  Sections: {len(resume['experience'])} experience entries, "
          f"{sum(len(v) for v in resume['notable_projects'].values())} projects")


if __name__ == "__main__":
    main()
